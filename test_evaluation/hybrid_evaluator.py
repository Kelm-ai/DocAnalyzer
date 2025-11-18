#!/usr/bin/env python3
"""
ISO 14971 Hybrid Evaluator

Combines the markdown-oriented flow with OpenAI's vision capabilities.
We supply the full markdown excerpt alongside the requirement prompt and
also attach the original document (PDF/DOCX) via the Files API so the
model can reference visuals when needed.
"""

from __future__ import annotations

import argparse
import asyncio
import hashlib
import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# Load environment variables from .env file
from dotenv import load_dotenv
load_dotenv()

import PyPDF2
from openai import AsyncOpenAI
from openpyxl import Workbook
from openpyxl.utils import get_column_letter

from evaluation_schema import RequirementEvaluationSchema

try:
    from docx import Document  # type: ignore
    DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover - optional dependency
    Document = None  # type: ignore
    DOCX_AVAILABLE = False


class HybridEvaluator:
    """Evaluate ISO requirements using markdown context plus file attachment."""

    BASE_INSTRUCTION = (
        "You are an ISO 14971:2019 compliance auditor. Review the markdown context "
        "provided below. If details appear incomplete (e.g., figures, tables, or "
        "formatting nuances), consult the attached source document to confirm.")

    RESPONSE_SCHEMA = (
        "Respond strictly with JSON using this schema:\n"
        "{\n"
        "  \"status\": \"PASS|FAIL|FLAGGED|NOT_APPLICABLE\",\n"
        "  \"confidence\": \"low|medium|high\",\n"
        "  \"rationale\": string,\n"
        "  \"evidence\": [string],\n"
        "  \"gaps\": [string],\n"
        "  \"recommendations\": [string]\n"
        "}\n"
        "Confidence level guidelines:\n"
        "- Use \"high\" when evidence is explicit, comprehensive, and directly addresses all criteria\n"
        "- Use \"medium\" when evidence is present but incomplete, requires some inference, or has minor gaps\n"
        "- Use \"low\" when evidence is sparse, ambiguous, uncertain, or requires significant assumptions\n"
    )

    def __init__(
        self,
        *,
        openai_api_key: Optional[str] = None,
        model: Optional[str] = None,
    ) -> None:
        api_key = openai_api_key or os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise RuntimeError("OPENAI_API_KEY is required for the hybrid evaluator")

        self.model = model or os.getenv("OPENAI_HYBRID_MODEL", os.getenv("OPENAI_MODEL", "gpt-5"))
        self.client = AsyncOpenAI(api_key=api_key)

        self.context_char_limit = int(os.getenv("HYBRID_CONTEXT_CHAR_LIMIT", "90000"))
        self.concurrent_requests = int(os.getenv("HYBRID_EVALUATOR_CONCURRENCY", "3"))
        self.reasoning_effort = os.getenv('HYBRID_REASONING_EFFORT', 'medium')

        base_dir = Path(__file__).parent
        self.base_dir = base_dir
        self.requirements_path = base_dir / "requirements_test.json"
        self.output_dir = base_dir / "output" / "hybrid_results"
        self.responses_dir = self.output_dir / "responses"
        self.markdown_dir = base_dir / "output" / "hybrid_markdown"

        for directory in (self.output_dir, self.responses_dir, self.markdown_dir):
            directory.mkdir(parents=True, exist_ok=True)

        # Shared cache so other evaluators can reuse uploads
        self.cache_path = base_dir / "output" / "uploaded_files_cache.json"
        self.file_cache = self._load_cache()

    async def evaluate_document(self, file_path: str) -> Dict:
        document_path = Path(file_path)
        if not document_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        markdown = await asyncio.to_thread(self._convert_to_markdown, document_path)
        truncated_markdown = markdown[: self.context_char_limit]

        file_id, file_hash, cache_hit = await self.ensure_file_id(document_path)
        requirements = self._load_requirements()

        run_id = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        run_responses_dir = self.responses_dir / run_id
        run_responses_dir.mkdir(parents=True, exist_ok=True)

        markdown_file = self.markdown_dir / f"{document_path.stem}_{run_id}.md"
        markdown_file.write_text(markdown, encoding="utf-8")

        document_stats = {
            "file_name": document_path.name,
            "file_path": str(document_path.resolve()),
            "file_size_bytes": document_path.stat().st_size,
            "uploaded_file_id": file_id,
            "content_hash": file_hash,
            "file_id_cache_hit": cache_hit,
            "model": self.model,
            "evaluated_at": datetime.utcnow().isoformat(),
            "run_id": run_id,
            "markdown_chars": len(markdown),
            "markdown_truncated_chars": len(truncated_markdown),
            "markdown_file": str(markdown_file.resolve()),
        }

        semaphore = asyncio.Semaphore(self.concurrent_requests)
        tasks = [
            self._evaluate_single_requirement(
                file_id=file_id,
                requirement=req,
                semaphore=semaphore,
                run_responses_dir=run_responses_dir,
                markdown_context=truncated_markdown,
            )
            for req in requirements
        ]

        evaluations = await asyncio.gather(*tasks, return_exceptions=True)

        results: List[Dict] = []
        for requirement, evaluation in zip(requirements, evaluations):
            if isinstance(evaluation, Exception):
                results.append({
                    "requirement_id": requirement["id"],
                    "status": "ERROR",
                    "confidence": "low",  # Categorical string confidence
                    "rationale": str(evaluation),
                    "evidence": [],
                    "gaps": ["Evaluation failed"],
                    "recommendations": ["Retry requirement"],
                    "tokens_used": 0,
                })
            else:
                results.append(evaluation)

        summary = self._generate_summary(document_stats, results)
        self._persist_summary(summary, run_id)
        return summary

    async def ensure_file_id(self, document_path: Path) -> Tuple[str, str, bool]:
        data = document_path.read_bytes()
        file_hash = hashlib.sha256(data).hexdigest()
        cached_entry = self.file_cache.get(file_hash)
        if cached_entry:
            return cached_entry["file_id"], file_hash, True

        with open(document_path, "rb") as file_obj:
            upload = await self.client.files.create(file=file_obj, purpose="user_data")

        self.file_cache[file_hash] = {
            "file_id": upload.id,
            "uploaded_at": datetime.utcnow().isoformat(),
            "file_name": document_path.name,
        }
        self._save_cache()
        return upload.id, file_hash, False

    async def _evaluate_single_requirement(
        self,
        *,
        file_id: str,
        requirement: Dict,
        semaphore: asyncio.Semaphore,
        run_responses_dir: Path,
        markdown_context: str,
    ) -> Dict:
        async with semaphore:
            prompt = self._build_prompt(requirement, markdown_context)

            response = await self.client.responses.parse(
                model=self.model,
                reasoning={"effort": self.reasoning_effort},
                input=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "input_text", "text": prompt},
                            {"type": "input_file", "file_id": file_id},
                        ],
                    }
                ],
                text_format=RequirementEvaluationSchema,
            )

            usage = getattr(response, "usage", None)
            tokens_used = getattr(usage, "total_tokens", 0) if usage else 0

            raw_file = run_responses_dir / f"response_{requirement['id'].replace('-', '_')}.txt"
            parsed_model = getattr(response, "output_parsed", None)
            if parsed_model is None:
                raw_file.write_text("", encoding="utf-8")
                return {
                    "requirement_id": requirement["id"],
                    "status": "ERROR",
                    "confidence": "low",  # Categorical string confidence
                    "rationale": "Structured output missing from model response",
                    "evidence": [],
                    "gaps": ["Model response missing structured payload"],
                    "recommendations": ["Retry evaluation"],
                    "tokens_used": tokens_used,
                }

            parsed = parsed_model.model_dump()
            parsed.setdefault("requirement_id", requirement["id"])
            parsed["tokens_used"] = tokens_used
            raw_text = getattr(response, "output_text", None) or json.dumps(parsed, indent=2)
            raw_file.write_text(raw_text, encoding="utf-8")
            return parsed

    def _build_prompt(self, requirement: Dict, markdown_context: str) -> str:
        sections = [self.BASE_INSTRUCTION, self.RESPONSE_SCHEMA]

        sections.append(
            "COMBINED MARKDOWN CONTEXT (truncated):\n"
            f"{markdown_context}\n"
        )

        sections.append(
            "MANDATORY METHOD:\n"
            "1. Start with the markdown excerpt; consult the attached document when additional visual or formatting detail is needed.\n"
            "2. Evaluate each acceptance criterion individually and cite page or section references from either source.\n"
            "3. Use PASS when all criteria are clearly satisfied with explicit evidence, FAIL when evidence is clearly missing or contradictory, and FLAGGED only when the evidence is partial or genuinely uncertain.\n"
            "4. Confirm that the final status reflects the strength of the evidence; avoid defaulting to FLAGGED when PASS or FAIL is clearly supported.\n"
            "Respond strictly with JSON using this schema:\n"
            "{\n"
            "  \"status\": \"PASS|FAIL|FLAGGED|NOT_APPLICABLE\",\n"
            "  \"confidence\": \"low|medium|high\",\n"
            "  \"rationale\": \"Explain satisfied/unsatisfied criteria with citations\",\n"
            "  \"evidence\": [\"Page/Section citation with quote\", ...],\n"
            "  \"gaps\": [string],\n"
            "  \"recommendations\": [string]\n"
            "}\n"
            "Confidence level guidelines:\n"
            "- Use \"high\" when evidence is explicit, comprehensive, and directly addresses all criteria\n"
            "- Use \"medium\" when evidence is present but incomplete, requires some inference, or has minor gaps\n"
            "- Use \"low\" when evidence is sparse, ambiguous, uncertain, or requires significant assumptions\n"
        )

        sections.append(
            "Requirement to evaluate:\n"
            f"- ID: {requirement['id']}\n"
            f"- Clause: {requirement['clause']}\n"
            f"- Title: {requirement['title']}\n"
            f"- Requirement Text: {requirement['requirement_text']}\n"
            f"- Acceptance Criteria: {requirement['acceptance_criteria']}\n"
            f"- Expected Artifacts: {requirement.get('expected_artifacts', 'Not specified')}"
        )

        sections.append(
            "When referencing evidence, include short quotes with page/section identifiers. If the excerpt omits critical details, state that you confirmed them in the attachment.\n"
        )


        return "\n\n".join(sections)

    def _convert_to_markdown(self, document_path: Path) -> str:
        suffix = document_path.suffix.lower()
        if suffix == ".pdf":
            return self._convert_pdf_to_markdown(document_path)
        if suffix in {".docx", ".doc"}:
            if not DOCX_AVAILABLE:
                raise RuntimeError("python-docx is required to process DOC/DOCX files")
            return self._convert_docx_to_markdown(document_path)
        raise ValueError(f"Unsupported file type: {suffix}")

    def _convert_pdf_to_markdown(self, document_path: Path) -> str:
        markdown_sections: List[str] = []
        with open(document_path, "rb") as fh:
            reader = PyPDF2.PdfReader(fh)
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text() or ""
                normalized = self._normalize_pdf_text(text)
                if normalized.strip():
                    markdown_sections.append(f"## Page {page_num}\n\n{normalized}\n")
        markdown = "\n".join(markdown_sections).strip()
        if not markdown:
            raise ValueError("No extractable text found in PDF; markdown context unavailable")
        return markdown

    def _convert_docx_to_markdown(self, document_path: Path) -> str:
        doc = Document(str(document_path))  # type: ignore[arg-type]
        sections: List[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if paragraph.style and paragraph.style.name.startswith("Heading"):
                try:
                    level = int(re.findall(r"\d+", paragraph.style.name)[0])
                except (IndexError, ValueError):
                    level = 1
                level = max(1, min(level, 6))
                sections.append(f"{'#' * level} {text}")
            else:
                sections.append(text)

        for index, table in enumerate(doc.tables, start=1):
            sections.append(f"\n## Table {index}")
            header_written = False
            for row in table.rows:
                cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
                line = "| " + " | ".join(cells) + " |"
                sections.append(line)
                if not header_written and cells:
                    sections.append("| " + " | ".join(["---"] * len(cells)) + " |")
                    header_written = True
            sections.append("")

        markdown = "\n".join(sections).strip()
        if not markdown:
            raise ValueError("No extractable text found in DOC/DOCX document")
        return markdown

    def _normalize_pdf_text(self, text: str) -> str:
        lines = text.splitlines()
        markdown_lines: List[str] = []
        for line in lines:
            cleaned = " ".join(line.split())
            if not cleaned:
                continue
            if len(cleaned) < 100 and cleaned.isupper() and len(cleaned.split()) <= 10:
                markdown_lines.append(f"### {cleaned.title()}")
            else:
                markdown_lines.append(cleaned)
        return "\n\n".join(markdown_lines)

    def _extract_response_text(self, response) -> str:
        try:
            output = getattr(response, "output", None)
            if not output:
                return ""
            parts: List[str] = []
            for item in output:
                content = getattr(item, "content", None)
                if not content:
                    continue
                for chunk in content:
                    chunk_type = getattr(chunk, "type", None)
                    if chunk_type in {"output_text", "text"}:
                        parts.append(getattr(chunk, "text", ""))
            return "\n".join(filter(None, parts))
        except AttributeError:
            return ""

    def _load_requirements(self) -> List[Dict]:
        return json.loads(self.requirements_path.read_text())

    def _generate_summary(self, document_stats: Dict, results: List[Dict]) -> Dict:
        status_counts: Dict[str, int] = {"PASS": 0, "FAIL": 0, "FLAGGED": 0, "NOT_APPLICABLE": 0, "ERROR": 0}
        total_tokens = 0
        for result in results:
            status = result.get("status", "ERROR")
            status_counts[status] = status_counts.get(status, 0) + 1
            total_tokens += result.get("tokens_used", 0)

        scored = len(results) - status_counts.get("ERROR", 0)
        compliance_score = (status_counts.get("PASS", 0) / scored * 100) if scored else 0

        summary = {
            "document_info": document_stats,
            "evaluation_summary": {
                "total_requirements": len(results),
                "compliance_score": round(compliance_score, 1),
                "status_counts": status_counts,
                "total_tokens_used": total_tokens,
                "estimated_cost_usd": round((total_tokens / 1_000_000) * 5, 4),
            },
            "requirements_results": results,
            "generated_at": datetime.utcnow().isoformat(),
        }
        return summary

    def _persist_summary(self, summary: Dict, run_id: str) -> None:
        json_path = self.output_dir / f"hybrid_evaluation_{run_id}.json"
        json_path.write_text(json.dumps(summary, indent=2), encoding="utf-8")

        excel_path = self.output_dir / f"hybrid_evaluation_{run_id}.xlsx"
        self._export_to_excel(summary, excel_path)

    def _export_to_excel(self, summary: Dict, excel_path: Path) -> None:
        workbook = Workbook()
        summary_sheet = workbook.active
        summary_sheet.title = "Summary"

        summary_sheet.append(["Field", "Value"])
        for key, value in summary.get("document_info", {}).items():
            summary_sheet.append([key.replace('_', ' ').title(), value])

        summary_sheet.append([])
        summary_sheet.append(["Metric", "Value"])
        evaluation_summary = summary.get("evaluation_summary", {})
        for key, value in evaluation_summary.items():
            if key == "status_counts":
                continue
            summary_sheet.append([key.replace('_', ' ').title(), value])

        summary_sheet.append([])
        summary_sheet.append(["Status", "Count"])
        for status, count in evaluation_summary.get("status_counts", {}).items():
            summary_sheet.append([status, count])

        self._auto_size_columns(summary_sheet)

        requirements_sheet = workbook.create_sheet(title="Requirements")
        headers = [
            "Requirement ID",
            "Status",
            "Confidence",
            "Rationale",
            "Evidence",
            "Gaps",
            "Recommendations",
            "Tokens Used",
        ]
        requirements_sheet.append(headers)

        for record in summary.get("requirements_results", []):
            # Normalize confidence to uppercase categorical label
            confidence_raw = record.get("confidence", "low")
            confidence_str = str(confidence_raw).strip().lower()
            if confidence_str not in ("low", "medium", "high"):
                confidence_str = "low"
            confidence_label = confidence_str.upper()

            requirements_sheet.append([
                record.get("requirement_id"),
                record.get("status"),
                confidence_label,
                record.get("rationale", ""),
                "\n".join(record.get("evidence", [])),
                "\n".join(record.get("gaps", [])),
                "\n".join(record.get("recommendations", [])),
                record.get("tokens_used", 0),
            ])

        self._auto_size_columns(requirements_sheet)
        workbook.save(excel_path)

    def _auto_size_columns(self, worksheet) -> None:
        for column_cells in worksheet.columns:
            length = max(len(str(cell.value or "")) for cell in column_cells)
            worksheet.column_dimensions[get_column_letter(column_cells[0].column)].width = min(length + 2, 80)

    def _load_cache(self) -> Dict[str, Dict[str, str]]:
        if self.cache_path.exists():
            try:
                return json.loads(self.cache_path.read_text())
            except json.JSONDecodeError:
                return {}
        return {}

    def _save_cache(self) -> None:
        self.cache_path.write_text(json.dumps(self.file_cache, indent=2), encoding="utf-8")


async def _async_main(file_path: str) -> None:
    evaluator = HybridEvaluator()
    summary = await evaluator.evaluate_document(file_path)

    counts = summary["evaluation_summary"]["status_counts"]
    print("\n=== Hybrid Evaluation Complete ===")
    print(f"Document: {summary['document_info']['file_name']}")
    print(f"Model: {summary['document_info']['model']}")
    print(f"Score: {summary['evaluation_summary']['compliance_score']:.1f}%")
    print(f"Tokens used: {summary['evaluation_summary']['total_tokens_used']}")
    print("Status counts:")
    for status, count in counts.items():
        print(f"  {status:<15} {count}")


def main() -> None:
    parser = argparse.ArgumentParser(description="ISO 14971 hybrid evaluator")
    parser.add_argument("file_path", help="Path to PDF or DOCX file to evaluate")
    args = parser.parse_args()

    asyncio.run(_async_main(args.file_path))


if __name__ == "__main__":
    main()
