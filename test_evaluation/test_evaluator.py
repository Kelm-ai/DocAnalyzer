#!/usr/bin/env python3
"""
ISO 14971 Test Evaluator
Test version for document processing and evaluation with first 3 requirements
"""

import os
import json
import sys
import time
import argparse
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Optional

# Load environment variables from .env file
from dotenv import load_dotenv
load_dotenv()

from openpyxl import Workbook
from openpyxl.utils import get_column_letter

# External libraries
import PyPDF2
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX files won't be supported.")

try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False
    print("Warning: openai not installed. Evaluation won't work.")

from evaluation_schema import RequirementEvaluationSchema

try:
    from colorama import init, Fore, Style
    init()  # Initialize colorama
    COLORAMA_AVAILABLE = True
except ImportError:
    COLORAMA_AVAILABLE = False
    # Fallback color functions
    class Fore:
        RED = GREEN = YELLOW = BLUE = CYAN = MAGENTA = WHITE = RESET = ""
    class Style:
        BRIGHT = DIM = RESET_ALL = ""

try:
    from tabulate import tabulate
    TABULATE_AVAILABLE = True
except ImportError:
    TABULATE_AVAILABLE = False


class TestEvaluator:
    """Test version of ISO 14971 evaluator with markdown output and limited requirements"""

    def __init__(self, openai_api_key: str = None):
        self.openai_api_key = openai_api_key or os.getenv('OPENAI_API_KEY')
        if not self.openai_api_key:
            print(f"{Fore.RED}Error: OPENAI_API_KEY not found in environment{Style.RESET_ALL}")
            sys.exit(1)

        self.client = OpenAI(api_key=self.openai_api_key) if OPENAI_AVAILABLE else None
        self.model = os.getenv('OPENAI_MODEL', 'gpt-5')
        self.document_context_char_limit = int(os.getenv('DOCUMENT_CONTEXT_CHAR_LIMIT', '90000'))
        self.reasoning_effort = os.getenv('EVALUATOR_REASONING_EFFORT', 'medium')

        # Setup output directories
        self.base_dir = Path(__file__).parent
        self.output_dir = self.base_dir / "output"
        self.markdown_dir = self.output_dir / "markdown"
        self.results_dir = self.output_dir / "results"
        self.reports_dir = self.output_dir / "reports"

        # Ensure directories exist
        for dir_path in [self.markdown_dir, self.results_dir, self.reports_dir]:
            dir_path.mkdir(parents=True, exist_ok=True)

    def print_header(self, text: str, color=Fore.CYAN):
        """Print a colored header"""
        print(f"\n{color}{'='*60}")
        print(f"{text:^60}")
        print(f"{'='*60}{Style.RESET_ALL}\n")

    def print_status(self, text: str, status: str = "INFO"):
        """Print colored status message"""
        colors = {
            "INFO": Fore.BLUE,
            "SUCCESS": Fore.GREEN,
            "WARNING": Fore.YELLOW,
            "ERROR": Fore.RED,
            "PROCESSING": Fore.MAGENTA
        }
        color = colors.get(status, Fore.WHITE)
        print(f"{color}[{status}]{Style.RESET_ALL} {text}")

    def convert_pdf_to_markdown(self, file_path: str) -> str:
        """Convert PDF to markdown format"""
        self.print_status("Converting PDF to markdown...", "PROCESSING")
        markdown_sections = []

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text() or ""
                    normalized = self._normalize_pdf_text(page_text)
                    if normalized.strip():
                        markdown_sections.append(f"## Page {page_num}\n\n{normalized}\n")

                markdown = '\n'.join(markdown_sections).strip()

                if not markdown:
                    raise ValueError("No extractable text found in PDF")

                self.print_status(f"PDF converted: {len(pdf_reader.pages)} pages, {len(markdown)} characters", "SUCCESS")
                return markdown

        except Exception as e:
            self.print_status(f"PDF conversion failed: {e}", "ERROR")
            raise

    def convert_docx_to_markdown(self, file_path: str) -> str:
        """Convert DOCX to markdown format"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available for DOCX processing")

        self.print_status("Converting DOCX to markdown...", "PROCESSING")
        markdown_sections = []

        try:
            doc = Document(file_path)

            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                text = paragraph.text.strip()
                if text:
                    # Check if it's a heading (basic detection)
                    if paragraph.style.name.startswith('Heading'):
                        level = min(int(paragraph.style.name.split()[-1]), 6)
                        markdown_sections.append(f"{'#' * level} {text}\n")
                    else:
                        markdown_sections.append(f"{text}\n")

            # Process tables
            for table_num, table in enumerate(doc.tables, 1):
                markdown_sections.append(f"\n## Table {table_num}\n")
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    markdown_sections.append(f"| {row_text} |")
                markdown_sections.append("")

            markdown = '\n'.join(markdown_sections).strip()

            if not markdown:
                raise ValueError("No extractable text found in DOCX")

            self.print_status(f"DOCX converted: {len(doc.paragraphs)} paragraphs, {len(markdown)} characters", "SUCCESS")
            return markdown

        except Exception as e:
            self.print_status(f"DOCX conversion failed: {e}", "ERROR")
            raise

    def _normalize_pdf_text(self, text: str) -> str:
        """Normalize extracted PDF text"""
        import re

        # Split into lines and clean
        lines = text.split('\n')
        markdown_lines = []

        for line in lines:
            # Remove excessive whitespace
            cleaned = ' '.join(line.split())

            # Skip empty lines
            if not cleaned:
                continue

            # Basic heading detection (all caps, short lines)
            if len(cleaned) < 100 and cleaned.isupper() and len(cleaned.split()) <= 10:
                markdown_lines.append(f"### {cleaned.title()}")
            else:
                markdown_lines.append(cleaned)

        # Join with proper spacing
        result = '\n\n'.join(markdown_lines)

        # Clean up excessive newlines
        result = re.sub(r'\n{3,}', '\n\n', result)

        return result

    def convert_document_to_markdown(self, file_path: str) -> tuple[str, dict]:
        """Convert document to markdown and return content + stats"""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Detect file type and convert
        suffix = file_path.suffix.lower()

        if suffix == '.pdf':
            markdown = self.convert_pdf_to_markdown(str(file_path))
        elif suffix in ['.docx', '.doc']:
            markdown = self.convert_docx_to_markdown(str(file_path))
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        # Calculate statistics
        stats = {
            "file_name": file_path.name,
            "file_type": suffix,
            "file_size_bytes": file_path.stat().st_size,
            "markdown_length": len(markdown),
            "word_count": len(markdown.split()),
            "line_count": len(markdown.split('\n')),
            "conversion_timestamp": datetime.now().isoformat()
        }

        # Save markdown to file
        markdown_file = self.markdown_dir / f"{file_path.stem}.md"
        markdown_file.write_text(markdown, encoding='utf-8')
        self.print_status(f"Markdown saved to: {markdown_file}", "SUCCESS")

        return markdown, stats

    def load_test_requirements(self) -> List[Dict]:
        """Load the 3 test requirements"""
        requirements_file = self.base_dir / "requirements_test.json"

        try:
            with open(requirements_file, 'r') as f:
                requirements = json.load(f)

            self.print_status(f"Loaded {len(requirements)} test requirements", "SUCCESS")
            return requirements

        except Exception as e:
            self.print_status(f"Failed to load requirements: {e}", "ERROR")
            raise

    def evaluate_single_requirement(self, document_markdown: str, requirement: Dict) -> Dict:
        """Evaluate a single requirement against the document"""
        if not OPENAI_AVAILABLE or not self.client:
            self.print_status("OpenAI not available - skipping evaluation", "WARNING")
            return {
                "requirement_id": requirement['id'],
                "status": "SKIPPED",
                "confidence": "low",
                "rationale": "OpenAI API not available",
                "evidence": [],
                "gaps": ["OpenAI API not configured"],
                "recommendations": ["Configure OpenAI API key"],
                "tokens_used": 0,
                "evaluation_duration_ms": 0
            }

        start_time = time.time()

        # Truncate document if too long
        context_snippet = document_markdown[:self.document_context_char_limit]
        if len(document_markdown) > self.document_context_char_limit:
            self.print_status(f"Document truncated to {self.document_context_char_limit} characters", "WARNING")

        # Build evaluation prompt
        prompt = f"""You are an ISO 14971:2019 compliance auditor. Review the markdown context below first. If visuals or formatting details are unclear, you may rely on the original document as needed when forming your judgement.

MANDATORY METHOD:
1. Examine each acceptance criterion individually and explain in your rationale whether it is satisfied.
2. Provide explicit evidence with page or section references (e.g., "Page 4: ...").
3. Output PASS when every criterion is clearly satisfied with cited evidence. Use FAIL when evidence is clearly missing or contradictory. Reserve FLAGGED for cases where evidence is partial or genuinely uncertain.
4. Before finalising, confirm that the chosen status (PASS / FAIL / FLAGGED) best reflects the evidence; do not default to FLAGGED when the evidence clearly supports PASS or FAIL.

MARKDOWN CONTEXT (truncated to {self.document_context_char_limit} chars):
{context_snippet}

REQUIREMENT DETAILS:
- ID: {requirement['id']}
- Clause: {requirement['clause']}
- Title: {requirement['title']}
- Requirement Text: {requirement['requirement_text']}
- Acceptance Criteria: {requirement['acceptance_criteria']}
- Expected Artifacts: {requirement.get('expected_artifacts', 'Not specified')}

Respond with JSON only:
{{
    "status": "PASS|FAIL|FLAGGED|NOT_APPLICABLE",
    "confidence": "low|medium|high",
    "rationale": "Explain satisfied/unsatisfied criteria with citations",
    "evidence": ["Page/Section citation with quote", ...],
    "gaps": ["Gap 1", ...],
    "recommendations": ["Next action", ...]
}}
Confidence level guidelines:
- Use "high" when evidence is explicit, comprehensive, and directly addresses all criteria
- Use "medium" when evidence is present but incomplete, requires some inference, or has minor gaps
- Use "low" when evidence is sparse, ambiguous, uncertain, or requires significant assumptions"""

        self.print_status(f"Evaluating requirement {requirement['id']}...", "PROCESSING")

        # Save prompt for debugging
        prompt_file = self.results_dir / f"prompt_{requirement['id'].replace('-', '_')}.txt"
        prompt_file.write_text(prompt, encoding='utf-8')

        try:
            # Some models like gpt-5-mini don't support custom temperature
            # Use the Responses API so reasoning parameters are supported consistently
            # Note: gpt-5 models don't support temperature parameter with responses API

            response = self.client.responses.parse(
                model=self.model,
                reasoning={"effort": self.reasoning_effort},
                input=[{
                    "role": "user",
                    "content": [{"type": "input_text", "text": prompt}]
                }],
                text_format=RequirementEvaluationSchema,
            )

            tokens_used = getattr(getattr(response, "usage", None), "total_tokens", 0)

            # Save raw response
            response_file = self.results_dir / f"response_{requirement['id'].replace('-', '_')}.txt"
            parsed_model = getattr(response, "output_parsed", None)
            if parsed_model is None:
                response_file.write_text("", encoding='utf-8')
                return {
                    "requirement_id": requirement['id'],
                    "status": "ERROR",
                    "confidence": "low",  # Categorical string confidence
                    "rationale": "Structured output missing from model response",
                    "evidence": [],
                    "gaps": ["Model response missing structured payload"],
                    "recommendations": ["Retry evaluation"],
                    "tokens_used": tokens_used,
                    "evaluation_duration_ms": int((time.time() - start_time) * 1000),
                }

            parsed = parsed_model.model_dump()
            parsed['requirement_id'] = requirement['id']
            parsed['tokens_used'] = tokens_used
            parsed['evaluation_duration_ms'] = int((time.time() - start_time) * 1000)
            raw_text = getattr(response, "output_text", None) or json.dumps(parsed, indent=2)
            response_file.write_text(raw_text, encoding='utf-8')

            # Display categorical confidence level
            confidence_display = str(parsed.get('confidence', 'low')).upper()
            self.print_status(
                f"Evaluation complete: {parsed['status']} (confidence: {confidence_display})",
                "SUCCESS",
            )
            return parsed

        except Exception as e:
            self.print_status(f"API call failed: {e}", "ERROR")
            return {
                "requirement_id": requirement['id'],
                "status": "ERROR",
                "confidence": "low",  # Categorical string confidence
                "rationale": f"API error: {e}",
                "evidence": [],
                "gaps": ["API call failed"],
                "recommendations": ["Check API configuration"],
                "tokens_used": 0,
                "evaluation_duration_ms": int((time.time() - start_time) * 1000)
            }

    def generate_summary_report(self, document_stats: dict, results: List[Dict]) -> Dict:
        """Generate summary report"""
        total_requirements = len(results)

        status_counts = {
            "PASS": 0,
            "FAIL": 0,
            "FLAGGED": 0,
            "NOT_APPLICABLE": 0,
            "ERROR": 0,
            "SKIPPED": 0
        }

        total_tokens = 0
        total_duration = 0

        for result in results:
            status = result.get('status', 'ERROR')
            status_counts[status] = status_counts.get(status, 0) + 1
            total_tokens += result.get('tokens_used', 0)
            total_duration += result.get('evaluation_duration_ms', 0)

        # Calculate compliance score
        scored_requirements = total_requirements - status_counts['ERROR'] - status_counts['SKIPPED']
        if scored_requirements > 0:
            compliance_score = (status_counts['PASS'] / scored_requirements) * 100
        else:
            compliance_score = 0

        estimated_cost = (total_tokens / 1_000_000) * 5  # ~$5 per 1M tokens

        summary = {
            "document_info": document_stats,
            "evaluation_summary": {
                "total_requirements": total_requirements,
                "compliance_score": round(compliance_score, 1),
                "status_counts": status_counts,
                "total_tokens_used": total_tokens,
                "total_duration_ms": total_duration,
                "estimated_cost_usd": round(estimated_cost, 4)
            },
            "requirements_results": results,
            "generated_at": datetime.now().isoformat()
        }

        return summary

    def print_results_table(self, results: List[Dict]):
        """Print results in a formatted table"""
        if not TABULATE_AVAILABLE:
            # Fallback simple table
            print(f"\n{Fore.CYAN}EVALUATION RESULTS{Style.RESET_ALL}")
            print("-" * 80)
            for result in results:
                status_color = {
                    "PASS": Fore.GREEN,
                    "FAIL": Fore.RED,
                    "FLAGGED": Fore.YELLOW,
                    "NOT_APPLICABLE": Fore.BLUE,
                    "ERROR": Fore.RED,
                    "SKIPPED": Fore.MAGENTA
                }.get(result['status'], Fore.WHITE)

                confidence_str = str(result.get('confidence', 'low')).upper()
                print(f"{result['requirement_id']:20} {status_color}{result['status']:12}{Style.RESET_ALL} {confidence_str:8} {result.get('tokens_used', 0):8}")
            return

        # Use tabulate for nice formatting
        table_data = []
        for result in results:
            confidence_str = str(result.get('confidence', 'low')).upper()
            table_data.append([
                result['requirement_id'],
                result['status'],
                confidence_str,
                result.get('tokens_used', 0),
                f"{result.get('evaluation_duration_ms', 0)}ms"
            ])

        headers = ["Requirement ID", "Status", "Confidence", "Tokens", "Duration"]
        print(f"\n{Fore.CYAN}EVALUATION RESULTS{Style.RESET_ALL}")
        print(tabulate(table_data, headers=headers, tablefmt="grid"))

    def run_evaluation(self, file_path: str) -> Dict:
        """Run complete evaluation process"""
        self.print_header("ISO 14971 Test Evaluator", Fore.CYAN)

        # Step 1: Convert document to markdown
        self.print_header("Step 1: Document Processing", Fore.BLUE)
        markdown, document_stats = self.convert_document_to_markdown(file_path)

        # Print markdown preview
        print(f"{Fore.YELLOW}Markdown Preview (first 500 chars):{Style.RESET_ALL}")
        print("-" * 60)
        print(markdown[:500] + "..." if len(markdown) > 500 else markdown)
        print("-" * 60)

        # Step 2: Load requirements
        self.print_header("Step 2: Loading Requirements", Fore.BLUE)
        requirements = self.load_test_requirements()

        # Step 3: Run evaluations
        self.print_header("Step 3: Running Evaluations", Fore.BLUE)
        results = []

        for i, requirement in enumerate(requirements, 1):
            print(f"\n{Fore.CYAN}[{i}/{len(requirements)}] {requirement['title']}{Style.RESET_ALL}")
            result = self.evaluate_single_requirement(markdown, requirement)
            results.append(result)

        # Step 4: Generate summary
        self.print_header("Step 4: Summary Report", Fore.BLUE)
        summary = self.generate_summary_report(document_stats, results)

        # Print results table
        self.print_results_table(results)

        # Print summary stats
        eval_summary = summary['evaluation_summary']
        print(f"\n{Fore.GREEN}SUMMARY STATISTICS{Style.RESET_ALL}")
        print(f"Total Requirements: {eval_summary['total_requirements']}")
        print(f"Compliance Score: {eval_summary['compliance_score']:.1f}%")
        print(f"Total Tokens Used: {eval_summary['total_tokens_used']:,}")
        print(f"Estimated Cost: ${eval_summary['estimated_cost_usd']:.4f}")
        print(f"Total Duration: {eval_summary['total_duration_ms']/1000:.1f} seconds")

        # Save complete results
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        results_file = self.results_dir / f"evaluation_{timestamp}.json"
        with open(results_file, 'w') as f:
            json.dump(summary, f, indent=2)

        self.print_status(f"Complete results saved to: {results_file}", "SUCCESS")

        # Save Excel summary
        excel_file = self.results_dir / f"evaluation_{timestamp}.xlsx"
        self.export_results_to_excel(summary, excel_file)
        self.print_status(f"Excel report saved to: {excel_file}", "SUCCESS")

        return summary

    def _extract_response_text(self, response) -> str:
        try:
            output = getattr(response, "output", None)
            if not output:
                return ""
            parts: List[str] = []
            for item in output:
                content = getattr(item, "content", None)
                if not content:
                    continue
                for chunk in content:
                    chunk_type = getattr(chunk, "type", None)
                    if chunk_type in {"output_text", "text"}:
                        parts.append(getattr(chunk, "text", ""))
            return "\n".join(filter(None, parts))
        except AttributeError:
            return ""

    def export_results_to_excel(self, summary: Dict, excel_path: Path) -> Path:
        """Export evaluation summary and requirement details to an Excel workbook."""
        workbook = Workbook()

        # Summary sheet
        summary_sheet = workbook.active
        summary_sheet.title = "Summary"

        document_info = summary.get('document_info', {})
        evaluation_summary = summary.get('evaluation_summary', {})

        summary_sheet.append(["Field", "Value"])
        for key, value in document_info.items():
            summary_sheet.append([key.replace('_', ' ').title(), value])

        summary_sheet.append([])
        summary_sheet.append(["Metric", "Value"])
        for key, value in evaluation_summary.items():
            if key == 'status_counts':
                continue
            summary_sheet.append([key.replace('_', ' ').title(), value])

        status_counts = evaluation_summary.get('status_counts', {})
        if status_counts:
            summary_sheet.append([])
            summary_sheet.append(["Status", "Count"])
            for status, count in status_counts.items():
                summary_sheet.append([status, count])

        self._auto_size_columns(summary_sheet)

        # Requirements sheet
        requirements_sheet = workbook.create_sheet(title="Requirements")
        headers = [
            "Requirement ID",
            "Status",
            "Confidence",
            "Rationale",
            "Evidence",
            "Gaps",
            "Recommendations",
            "Tokens Used",
            "Duration (ms)"
        ]
        requirements_sheet.append(headers)

        for requirement in summary.get('requirements_results', []):
            evidence = '\n'.join(requirement.get('evidence', []))
            gaps = '\n'.join(requirement.get('gaps', []))
            recommendations = '\n'.join(requirement.get('recommendations', []))

            # Normalize confidence to uppercase categorical label
            confidence_raw = requirement.get("confidence", "low")
            confidence_str = str(confidence_raw).strip().lower()
            if confidence_str not in ("low", "medium", "high"):
                confidence_str = "low"
            confidence_label = confidence_str.upper()

            requirements_sheet.append([
                requirement.get('requirement_id'),
                requirement.get('status'),
                confidence_label,
                requirement.get('rationale', ''),
                evidence,
                gaps,
                recommendations,
                requirement.get('tokens_used', 0),
                requirement.get('evaluation_duration_ms', 0)
            ])

        self._auto_size_columns(requirements_sheet)

        workbook.save(excel_path)
        return excel_path

    def _auto_size_columns(self, worksheet) -> None:
        """Automatically size worksheet columns based on content length."""
        for column_cells in worksheet.columns:
            length = max(len(str(cell.value or "")) for cell in column_cells)
            adjusted_width = min(length + 2, 80)
            worksheet.column_dimensions[get_column_letter(column_cells[0].column)].width = adjusted_width


def main():
    parser = argparse.ArgumentParser(description="ISO 14971 Test Evaluator")
    parser.add_argument("file_path", help="Path to PDF or DOCX file to evaluate")
    parser.add_argument("--api-key", help="OpenAI API key (or set OPENAI_API_KEY env var)")

    args = parser.parse_args()

    try:
        evaluator = TestEvaluator(openai_api_key=args.api_key)
        results = evaluator.run_evaluation(args.file_path)

        print(f"\n{Fore.GREEN}✓ Evaluation completed successfully!{Style.RESET_ALL}")

    except Exception as e:
        print(f"\n{Fore.RED}✗ Evaluation failed: {e}{Style.RESET_ALL}")
        sys.exit(1)


if __name__ == "__main__":
    main()
